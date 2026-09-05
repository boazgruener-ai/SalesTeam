# -*- coding: utf-8 -*-
import re
import zipfile
import shutil
import os
import docx
from docx.shared import Pt

SALESTEAM_DIR = "C:/Users/41793/OneDrive/Documents/Personal/Claude Code/SalesTeam"


def fix_zoom(path):
    tmp = path + ".tmp"
    shutil.copy(path, tmp)
    with zipfile.ZipFile(tmp, "r") as zin:
        names = zin.namelist()
        settings = zin.read("word/settings.xml").decode("utf-8")
        if "<w:zoom" not in settings:
            m = re.search(r"(<w:settings[^>]*>)", settings)
            if m:
                settings = settings[: m.end()] + '<w:zoom w:percent="100"/>' + settings[m.end() :]
        data = {n: zin.read(n) for n in names}
        data["word/settings.xml"] = settings.encode("utf-8")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    os.remove(tmp)


# Splits a line of inline markdown into (text, bold, code) runs - handles **bold** and `code` spans.
TOKEN_RE = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")


def add_inline_runs(paragraph, text):
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(1) is not None:
            r = paragraph.add_run(m.group(1))
            r.bold = True
        else:
            r = paragraph.add_run(m.group(2))
            r.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


with open(SALESTEAM_DIR + "/RELEASE_NOTES.md", encoding="utf-8") as f:
    lines = f.read().splitlines()

doc = docx.Document()

for raw_line in lines:
    line = raw_line.rstrip()
    if not line or line == "---":
        continue
    if line.startswith("# "):
        p = doc.add_heading(level=1)
        add_inline_runs(p, line[2:])
    elif line.startswith("## "):
        p = doc.add_heading(level=2)
        add_inline_runs(p, line[3:])
    elif line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, line[2:])
    else:
        p = doc.add_paragraph()
        add_inline_runs(p, line)

for section in doc.sections:
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)

out_path = SALESTEAM_DIR + "/RELEASE_NOTES.docx"
doc.save(out_path)
fix_zoom(out_path)
print("Saved RELEASE_NOTES.docx")
