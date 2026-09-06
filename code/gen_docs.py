# -*- coding: utf-8 -*-
import re
import zipfile
import shutil
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SALESTEAM_DIR = "C:/Users/41793/OneDrive/Documents/Personal/Claude Code/SalesTeam"


def fix_zoom(path):
    """python-docx's blank template omits <w:zoom>, which some strict OOXML
    validators (and occasionally Word itself) flag. Patch it in post-save."""
    tmp = path + ".tmp"
    shutil.copy(path, tmp)
    with zipfile.ZipFile(tmp, "r") as zin:
        names = zin.namelist()
        settings = zin.read("word/settings.xml").decode("utf-8")
        if "<w:zoom" not in settings:
            settings = settings.replace(
                "<w:settings",
                '<w:settings',
                1,
            )
            # Insert a zoom element right after the opening <w:settings ...> tag.
            m = re.search(r"(<w:settings[^>]*>)", settings)
            if m:
                settings = settings[:m.end()] + '<w:zoom w:percent="100"/>' + settings[m.end():]
        data = {}
        for n in names:
            data[n] = zin.read(n)
        data["word/settings.xml"] = settings.encode("utf-8")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    import os
    os.remove(tmp)


def add_bullets(doc, items):
    for item in items:
        if isinstance(item, tuple):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(item[0])
            r.bold = True
            p.add_run(item[1])
        else:
            doc.add_paragraph(item, style="List Bullet")


def add_para_with_bold_lead(doc, lead, rest, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(lead)
    r.bold = True
    p.add_run(rest)
    return p


# =====================================================================
# PRD.docx
# =====================================================================
doc = docx.Document()

doc.add_heading("SalesTeam — Product Requirements Document", level=1)

p = doc.add_paragraph()
r = p.add_run("Status: "); r.bold = True
p.add_run("Living document, reflects the shipped product as of v0.28.2.")
p = doc.add_paragraph()
r = p.add_run("Note: "); r.bold = True
p.add_run(
    "No PRD file existed for this project before this document — it was assembled now from the full build "
    "history to serve as the canonical, up-to-date spec going forward. Update it alongside future features "
    "rather than letting it drift from RELEASE_NOTES.md."
)

doc.add_heading("1. Problem", level=2)
doc.add_paragraph(
    "A B2B software/AI-services salesperson finds new leads by manually searching LinkedIn Posts and Jobs "
    "for relevant activity (hiring signals, AI-adoption posts, etc.), then has to separately figure out who's "
    "worth approaching, draft an opening message, and keep track of who they've already contacted — all by "
    "hand, with no tooling built for this specific workflow. Existing CRM/sales-intelligence tools assume "
    "leads already exist in a system; they don't help find them on LinkedIn in the first place."
)

doc.add_heading("2. Goals", level=2)
add_bullets(doc, [
    "Turn a repeatable LinkedIn search into a one-click, multi-topic scan across both Posts and Jobs.",
    "Surface only real, addressable signal — automatically filter out competitors and recruiter/staffing "
    "noise that matches the same keywords but is never a prospect.",
    "Prioritize the result automatically, so the salesperson always knows what to work on first without "
    "having to ask an AI mentor the same question after every scan.",
    "Give the salesperson an AI team on top of the data: a mentor for strategy, a simulated buyer to "
    "pressure-test messages against, and reviewed (never auto-sent) drafting.",
    "Do all of this safely: manual-trigger only (no scheduled/background automation), nothing sent without "
    "explicit review, and destructive actions (bulk status changes) require deliberate friction.",
])

doc.add_heading("3. Non-goals", level=2)
add_bullets(doc, [
    "Not a full CRM. No pipeline stages beyond a simple status label, no deal value/forecasting, no team "
    "features (single-user, local-only storage).",
    "Never sends a message on the user's behalf. Drafts are generated for copy-paste only.",
    "No scheduled or background scanning. Every scan is a manual click, by design — this keeps the tool's "
    "behavior indistinguishable from a careful human user, not an automation bot, for LinkedIn ToS reasons.",
    "Not a general LinkedIn scraper — scoped to the Posts and Jobs search-results pages the user is already "
    "viewing, using their own authenticated session.",
])

doc.add_heading("4. User", level=2)
doc.add_paragraph(
    "Grounded in a real B2B sales role in Swiss enterprise software/AI consulting: manually searching "
    "LinkedIn for hiring/AI-adoption signals, contacting people who match, and needing a lightweight way to "
    "track who's been approached without adopting a heavyweight CRM for a one-person prospecting workflow."
)

doc.add_heading("5. Architecture", level=2)
p = doc.add_paragraph()
r = p.add_run("Project layout (v0.27.0): "); r.bold = True
p.add_run(
    "the repo root holds only living project documentation (this PRD, RELEASE_NOTES.md, README.md) and "
    "unrelated assets (pitch deck, screenshots) - everything else is organized into subfolders:"
)
add_bullets(doc, [
    ("/code — ", "the actual loadable extension (manifest.json + every .js/.html/.css file + icons/) plus "
     "the Python doc-generation scripts. Chrome's unpacked-extension loading requires manifest.json and "
     "everything it references to live together in one folder (no ../ escapes allowed), so this is the "
     "folder Chrome's “Load unpacked” points at."),
    ("/backup — ", "Settings/Leads export downloads (manual and automatic pre-scan) land here."),
    ("/exports — ", "CSV export downloads land here, plus target-accounts.json (see 6.11), generated by "
     "code/convert_target_accounts.py."),
    ("/log — ", "periodic Activity Log file exports land here (see 6.10)."),
    ("/builds — ", "every shipped release's zip is archived under its own builds/vX.Y.Z/ folder, alongside "
     "a copy of that exact version's manifest.json."),
])
doc.add_paragraph(
    "Chrome/Edge Manifest V3 extension, no server of its own. Four pages share one chrome.storage.local "
    "dataset:"
)
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Page"
hdr[1].text = "Purpose"
rows_data = [
    ("Scanner (side panel)", "Topics (positive search keywords) + Negative Topics (auto-filters), Job "
     "Search config, the scan trigger, and the raw results list."),
    ("Dashboard (tab)", "The pipeline view — pie charts, a sortable/filterable/paginated lead table, "
     "per-lead detail page, bulk actions."),
    ("Advisors (tab)", "Cross-lead Sales Mentor and Customer Voice agent chats, for strategy questions not "
     "tied to one specific lead."),
    ("Settings (tab)", "Company context, Anthropic API key, message templates, output language, value-add "
     "offers — anything general-purpose, non-lead-specific. Home for any future general settings too."),
]
for name, purpose in rows_data:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = purpose

doc.add_paragraph()
doc.add_paragraph(
    "background.js (service worker) orchestrates scans: sequentially runs each topic's search in one "
    "background tab, merges results, applies negative-topic filtering, optionally re-applies filters to "
    "existing leads, then runs the automatic prioritization pass — before ever reporting the scan complete."
)
doc.add_paragraph(
    "agent-shared.js is the shared AI engine (system prompts, tool definitions, the Anthropic fetch/tool-use "
    "loop) used identically by the Scanner-era code, Dashboard, and Advisors — one implementation, not a "
    "drifting copy per page."
)

doc.add_heading("6. Feature spec", level=2)

doc.add_heading("6.1 Topics (Scanner)", level=3)
add_bullets(doc, [
    "Named keyword groups, with an optional second \u201cAND with\u201d group (post must match one keyword "
    "from each group). Same shape reused for both Post topics and Job-specific topics.",
    "LinkedIn's own search-complexity limits are worked around automatically via query chunking (max 6 "
    "OR-terms per group) - the user can add however many keywords they want, and the "
    "cost is made visible rather than hidden: each topic shows its own live “N searches for this "
    "topic” hint, and a grand total across every enabled Post and Job topic (“Total: 10 "
    "searches this scan will run”) sits above the “Scan All "
    "Topics” button, turning orange past 30.",
    "AND-topics search additively, not multiplicatively (v0.25.0) - found necessary once a user's own topic "
    "redesign (a large bilingual AND-topic chasing a previously-empty topic) pushed one real scan to 68 "
    "sub-queries. Rather than combining a topic's two keyword groups into one LinkedIn query per pairing "
    "(which required a full cartesian product of concept-chunks x activity-chunks to cover every "
    "combination), the two groups now run as two independent, cheap LinkedIn searches (each a plain OR "
    "list), and the AND is applied client-side by intersecting the two raw result sets on the post's "
    "profileUrl (not its own key, which is usually snippet-derived and, confirmed live in v0.25.0, unstable "
    "for the same post across two independent searches - fixed in v0.25.1) - same logical AND "
    "(concept-AND-activity is still required to count as a match), additive cost "
    "instead of multiplicative (a 30x30 topic dropped from 48 sub-queries to 10). Completely invisible in "
    "the Topics UI - same keywords + andKeywords shape, same editing experience. To offset the fact that "
    "each phase now searches a broader, single-constraint corpus (versus LinkedIn doing the full "
    "intersection server-side before), these two phases scrape twice as deep (more scroll passes, not more "
    "LinkedIn requests) so a genuine double-match has a better chance of surviving the client-side join.",
    "Author-title filter (checked client-side against each post's visible headline, never sent to LinkedIn) "
    "and an \u201cinclude in-post job ads\u201d toggle.",
    "\u201cSuggest Lookalike Topics\u201d button - looks at the salesperson's own highest-priority (P1-P3) "
    "Post leads and asks the Sales Mentor what made them strong matches, then suggests new keywords for "
    "finding more like them. Post leads only, deliberately: Job leads carry no scraped body text (only "
    "title/company/location), so there's nothing to generalize from beyond the title itself, and this only "
    "ever writes into Post Topics anyway. Review-first like everything else here: each suggestion shows its "
    "reasoning and lets the user add it to an existing Topic, create a new one, or skip it - nothing is "
    "written until explicitly accepted.",
    "“Analyze Post Search Quality” button (Search Quality section, after Negative Topics) - the "
    "actual fix for the problem Lookalike Topics can't help with (it only has good examples once Posts "
    "already score well). Looks at every unactioned Post lead - “New” and “Irrelevant” "
    "- plus the current Topics and Negative Topics together, diagnoses what's likely limiting quality/volume, "
    "and proposes specific keyword changes across both - additions, removals, or a genuinely new "
    "Topic/Negative Topic. Including Irrelevant leads is deliberate: an over-aggressive Negative Topic "
    "keyword is invisible otherwise, and stats include an irrelevantByNegativeTopic breakdown plus each "
    "Irrelevant example's exact irrelevantReason, so a false positive (a generic platform mention, an "
    "in-house HR poster) is directly visible rather than assumed correct. Grounded in the configured company "
    "context so it can propose genuinely new keywords, not just react to mediocre examples - explicitly told "
    "that a Negative Topic can only ever reduce volume, never fix a shortage of good leads. Same "
    "review-first pattern: each suggestion can be redirected to a different existing topic or accepted as "
    "new, and a removal referencing a keyword that's already gone is filtered out rather than shown as a "
    "broken action. Still excludes Dismissed/Contacted/Responded/Converted - those are the salesperson's "
    "own decisions, not the system's.",
])

doc.add_heading("6.2 Negative Topics (Lead Filters)", level=3)
add_bullets(doc, [
    "Same shape and matching logic as a positive Topic (keywords OR-group, optional AND-with group), but "
    "checked client-side against an already-scraped lead's own text after the search comes back — never "
    "sent to LinkedIn as a query, and never removes a search result, only marks it.",
    "Each topic has an appliesTo scope: Post leads only, Job listings only, or both — because a signal "
    "that's noise on one vertical (a recruiter's own post) is completely normal on the other (a job ad "
    "naming an HR contact).",
    "Checked against a Post lead's snippet + headline + company (Job leads: title + company) - including "
    "company (v0.23.0) matters because it lets a negative topic target who the poster actually works for, a "
    "much more precise signal than a keyword that can also match a mere passing mention in the post body.",
    "Three built-in topics, seeded by the Sales Mentor's own review of a real scan, cannot be removed but "
    "every field (keywords, AND-group, scope, enabled) is fully editable: Competitor Blocklist (applies to "
    "both - named competing consulting/services firms, deliberately excluding generic cloud/AI platform "
    "vendors like Microsoft/Google/AWS/NVIDIA, which get mentioned constantly as mere tooling references), "
    "Recruiter/Staffing Headline Filter (Post leads only - can false-positive on an in-house HR/Talent-"
    "Acquisition person posting their own employer's real opening), and Known Recruiting Firms (applies to "
    "both, v0.23.0 - a curated list of known Swiss recruiting/staffing agencies matched against company, a "
    "more reliable alternative to guessing from headline text).",
    "The user can add unlimited custom negative topics for any other recurring noise. New built-in defaults "
    "only apply to fresh installs - an existing configuration is the user's own live data and isn't touched "
    "automatically when a default changes.",
    "A match sets the lead's status to Irrelevant and records both the topic and the specific keyword that "
    "matched (irrelevantReason), shown as a hover tooltip on the Dashboard's status pill — distinct from "
    "Dismissed, which is always the salesperson's own decision, never the system's.",
    "\u201cAlso re-apply these filters to existing leads on the next scan\u201d checkbox (Scanner tile, "
    "unchecked by default, not a saved setting) — when checked, the next scan also re-checks every "
    "currently-\u201cNew\u201d existing lead against the current negative topics. Never touches a lead "
    "already acted on.",
    "\u201cApply Negative Filters\u201d button (Scanner tile, next to the Negative Topics list) \u2014 "
    "re-checks every existing lead against whatever's currently configured, instantly, with no new scan "
    "needed. Fully bidirectional: a \u201cNew\u201d lead that now matches becomes Irrelevant, and an "
    "\u201cIrrelevant\u201d lead that no longer matches (because a keyword was edited or removed) reverts "
    "to \u201cNew\u201d \u2014 the on-next-scan checkbox above only ever caught the first direction. Reports "
    "exactly how many leads moved each way. Never touches a lead already acted on.",
    "A lead already marked Irrelevant, or missing a reason (predates this feature), gets its reason "
    "backfilled automatically, best-effort, the next time it's read.",
])

doc.add_heading("6.3 Lead data model & statuses", level=3)
doc.add_paragraph(
    "Every lead (Post or Job listing) carries: status, statusUpdatedAt, priority + priorityReason + "
    "priorityScoredAt, irrelevantReason (if applicable), company + companyExtractedAt (if applicable), plus "
    "source-specific fields (author/headline/snippet for Posts; title/company/location for Jobs), "
    "firstSeenAt/lastSeenAt/postedAt, and matchedTopics."
)
doc.add_paragraph(
    "Statuses: New (default) \u2192 Contacted (manual, or automatic the moment a drafted message is copied) "
    "\u2192 Responded / Converted (manual) or Dismissed (manual) or Irrelevant (automatic, negative-topic "
    "match)."
)
p = doc.add_paragraph()
r = p.add_run("Company. "); r.bold = True
p.add_run(
    "Job leads get a clean company field straight from the scrape; Post leads never had one before v0.17.0 - "
    "only a free-text headline (e.g. \u201cHead of AI for IT @ Azqore\u201d). After every scan, a batched AI "
    "call extracts a best-guess company for every Post lead still missing one (silently skipped with no API "
    "key, same pattern as prioritization), with a \U0001F3E2 \u201cAssign Company\u201d row action to search "
    "every company already seen (native browser autocomplete) or type a new one, or clear it. A manually-set "
    "or previously-extracted company is never touched again by the automatic pass - companyExtractedAt is "
    "only ever set for an AI guess, cleared on manual assignment, so a scan can never silently overwrite a "
    "human's correction."
)

doc.add_heading("6.4 Automatic lead prioritization", level=3)
add_bullets(doc, [
    "After every scan's searches finish and negative-topic filtering has run (including any opted-in "
    "re-apply pass), every lead still \u201cNew\u201d without a priority is sent to the Sales Mentor in one "
    "batch call (not per-lead), scored P1 (highest — drop everything, contact today) to P5 (lowest — "
    "unlikely fit, low urgency), each with a short reason, via a forced structured tool call (not free-text "
    "parsing, so the output is always well-formed).",
    "Scoring weighs real fit against the configured company context, seniority/decision power, and genuine "
    "urgency signals — explicitly not just topical keyword overlap. Sees a Post lead's company (v0.23.0, "
    "once extracted or assigned) and its isJobAd/isHiringPost flags, and is explicitly told not to penalize "
    "an in-post job ad just because the poster personally isn't senior - the poster is often HR or an "
    "unrelated employee sharing the opening, not the eventual contact, so what matters is the company-level "
    "signal; the next step is finding a better contact there, not necessarily messaging the poster.",
    "Treats Ideal Customer Profile fit as a secondary, moderating factor, not a pass/fail gate (v0.24.3 - a "
    "real bug found via user-supplied examples: an explicit AI-engineer/AI-developer hiring signal was being "
    "crushed to P4-5 purely for being outside Switzerland or not “enterprise,” even though "
    "isHiringPost/isJobAd and company were all being seen correctly). An on-topic need outside the exact ICP "
    "should typically land around 2-3; the bottom of the range (4-5) is reserved for no genuine buying "
    "signal at all, a clearly unrelated technical domain, or noise that should already have been filtered.",
    "Guards against the opposite failure too (v0.24.0): a post from someone with an impressive AI-sounding "
    "title at a company that clearly already runs AI at scale, that's really just industry commentary or "
    "thought leadership - reacting to AI news, sharing opinions/trends - with no expressed need, project, "
    "challenge, or hire of its own. Topical overlap and an impressive title aren't buying intent; the prompt "
    "now explicitly scores that pattern low (4-5) regardless of how senior or on-topic the poster looks.",
    "Runs automatically, with no button — visible in the side panel as \u201cprioritizing N new leads\u2026\u201d "
    "before \u201cScan complete.\u201d Silently skipped (never fails the scan) if no Anthropic API key is "
    "configured.",
    "Never re-scores an already-scored lead, or a lead that isn't \u201cNew.\u201d",
    "\u201cPrioritize Unscored Leads\u201d button (Dashboard) catches up anything the automatic pass never "
    "reached — leads that predate the feature, or a scan that ran with no API key — scoring every unscored "
    "\u201cNew\u201d lead across the entire list, not just what's currently filtered on screen. Target "

    "Account matches (6.11) are applied first, deterministically, before whatever's left goes to the AI.",
    "\u201cRe-score All Priorities\u201d button (Dashboard, v0.24.0) re-runs the Mentor on every already-"
    "scored \u201cNew\u201d lead too, not just unscored ones - lets a prompt fix, a new/changed Ideal "

    "Customer Profile, or a freshly imported/updated Target Accounts list (6.11) retroactively apply to leads "
    "scored before it existed, with no new scan required. Same manual-override protection "
    "as everywhere else: a lead whose priority was set by hand (no priorityScoredAt) is never touched or "
    "resent to the AI. Confirms before running, since it overwrites existing AI-assigned priorities. The "
    "completion message reports two numbers: how many leads were successfully re-scored, and how many of "
    "those actually ended up with a different priority than before.",
    "Chunked batch prioritization (v0.24.2, applies to both buttons above): a single prioritize_leads call is "
    "capped at 8192 output tokens, which a large-enough batch (each lead needs a priority plus a written "
    "reason) could exceed mid-generation - coming back truncated/invalid and silently scoring nothing, with "
    "no way to tell that apart from “genuinely found nothing to score.” Both buttons now send 20 "
    "leads per AI call and apply each chunk's results as soon as it completes (so a later chunk failing "
    "doesn't lose earlier progress), and the status text shows real counted progress as each chunk lands - "
    "e.g. “Re-scoring 20 of 45 leads with the Sales Mentor…” - rather than only an "
    "elapsed-time guess with no sense of how much is left. The count reflects the chunk currently in flight, "
    "not only completed work (v0.24.5) - announced as each chunk starts, so the status never sits at a "
    "misleading “0 of N” while the first chunk is already running.",
    "Correlated re-scoring: if a scan finds a new lead from the same person (Post leads, matched by profile "
    "URL) or same company (Job leads, matched by normalized company) as an existing \u201cNew\u201d lead "
    "that's already scored, both get re-scored together in the same batch - a second signal from the same "
    "account can change the right priority. Still only ever touches \u201cNew\u201d leads; anything already "
    "acted on is never re-scored.",
    "Target Account matches (v0.28.0) are folded in before every batch call runs - during a scan, and in both "
    "Dashboard buttons above (v0.28.2, partitionLeadsByTargetAccount in storage.js) - so importing or "
    "updating the Target Accounts list retroactively re-prioritizes already-scanned leads via “Re-score "
    "All Priorities” without needing a fresh scan. See 6.11 for the full deterministic-vs-signal split.",
])

doc.add_heading("6.5 Dashboard", level=3)
add_bullets(doc, [
    "Pie charts (last 7 days / 30 days / all time), bucketed by each lead's real (parsed) post date, one "
    "colored slice per status; clicking a slice/legend row filters the table to that status.",
    "Table: Post Date, First Scanned, Source, Matched Topic / Matched Keywords (v0.24.4 - every Topic this "
    "lead matched and the specific keyword(s) that triggered each, deduplicated across topics; the same data "
    "the CSV export already carried, now visible and filterable directly in the table), Title, Content "
    "(3-line clamp, click to expand), Creator "
    "(link), Company, Connection, Status (with Irrelevant-reason tooltip), Priority (P1\u2013P5 colored pill, "
    "tooltip shows the Mentor's reason), Last Activity, Actions (Open/Edit, Consult Mentor, Send Message, "
    "Assign Company, Dismiss). Every sortable column supports click-to-sort and an Excel-style per-column "
    "dropdown (sort asc/desc, free-text filter). Column widths are user-resizable and persisted.",
    "\u201cGroup by Company\u201d checkbox - Excel-style outline grouping inside this same table (not a "
    "separate view): a collapsible header row per company (name, lead count, expand/collapse caret) with "
    "its leads nested underneath; leads with no company yet collect into a trailing \u201cUnknown "
    "company\u201d group. Each company header has a \u201cGet Account Summary\u201d button - a one-shot "
    "AI synthesis across every lead seen at that account, cached per session so re-opening it doesn't "
    "re-call the AI. Forces \u201cAll\u201d leads per page while active, restoring the prior page size "
    "when turned off.",
    "\u201cShow Irrelevant (negative-filtered) leads\u201d checkbox, unchecked by default and persisted — "
    "the general \u201cAll statuses\u201d view excludes Irrelevant leads so the table isn't dominated by "
    "filtered-out noise; explicitly selecting \u201cIrrelevant\u201d from the Status filter always shows "
    "them regardless.",
    "Pagination (20/50/100/All per page, remembered), global search (title/content/creator), CSV export "
    "(all leads, or exactly what's currently filtered — both include Priority).",
    "Bulk Change: a small, deliberately unobtrusive button (pagination row, not the main controls) opens a "
    "modal dialog — a red warning naming exactly how many currently-filtered leads will be affected, no "
    "default status pre-selected, a confirmation prompt on top of that, and an \u201cUndo Last Bulk "
    "Change\u201d button in the same dialog that restores every affected lead to its exact prior status (one "
    "level of undo, persists across Dashboard sessions until superseded by another bulk change). Closeable "
    "via a title-bar-style \u2715.",
    "Detail page (per lead): full content, status control, Priority override (v0.23.0 - a dropdown right "
    "next to Status lets the salesperson correct a priority the Mentor got wrong, or set one on a lead that "
    "was never scored; permanently protected from both the automatic per-scan pass and correlated "
    "re-scoring, since both key off whether the AI itself scored the lead, not just whether a priority "
    "exists), Draft Message (template-based, AI-generated, "
    "copy-to-clipboard — copying auto-advances status New \u2192 Contacted), and a lead-scoped "
    "Consult Mentor chat (persisted per lead) with two quick-action buttons - “Buyer Summary” "
    "and “Conversation Starters” - that send the same canned request an equivalent typed message "
    "would, through the identical conversation/history/tools, so the salesperson doesn't have to type the "
    "same standard requests for every lead.",
])

doc.add_heading("6.6 Advisors (Sales Mentor & Customer Voice)", level=3)
add_bullets(doc, [
    "Sales Mentor: persona-configurable agent for cross-lead strategy questions (\u201cwhich lead should I "
    "prioritize,\u201d general sales process advice) and drafting. Uses list_leads/get_lead_details tools "
    "only when a question actually needs real data — list_leads excludes Irrelevant leads and includes both "
    "Post and Job leads (tagged type + hasIndividualContact), instructing the Mentor to handle each "
    "appropriately.",
    "Customer Voice: persona-configurable agent that roleplays a realistic buyer, grounding itself in a "
    "named lead's real content when one is referenced, or a general persona otherwise.",
    "Both share one tool-use engine (runAgentTurn in agent-shared.js): bounded timeouts (2.5 min per "
    "conversational turn, 50s per tool call) with a live ticking status (\u201cThinking\u2026 (Ns)\u201d) so "
    "a genuinely long analysis reads as progress, not a hang.",
    "A stray Enter press while a turn is still in flight is a no-op (matches the Send button's disabled "
    "state) — previously this could start a second concurrent turn and silently lose a message.",
])

doc.add_heading("6.7 Settings", level=3)
doc.add_paragraph(
    "Language (English/German), company context (\u201cWhat We Offer,\u201d used by every AI feature to "
    "reason about real fit), Ideal Customer Profile (v0.23.0 - who's specifically being targeted: size, "
    "geography, what they're investing in; deliberately a separate field from \u201cWhat We Offer\u201d since "
    "the product and the target customer are different concepts, even though every relevant AI feature reads "
    "both together - not used by Customer Voice, which has no reason to reason about who the seller "
    "targets), Anthropic API key, message templates (auto-picked per lead by connection status "
    "/ job-ad detection, or chosen manually), value-add offers (a fixed list the AI may mention, never "
    "invents). Target Accounts (v0.28.0 - imported list plus the auto-Priority-1 score threshold; see 6.11). "
    "Opened via its own blue button in the Scanner tile. The Advisors page reads these live (via "
    "chrome.storage.onChanged) rather than caching a stale copy, since editing now happens on a separate page."
)

doc.add_heading("6.8 Backup / portability", level=3)
doc.add_paragraph(
    "Two independent export/import flows (side panel, v0.25.1) - restoring one can never touch or roll back "
    "the other, found necessary after a combined single-file design meant recovering lost leads also "
    "silently rolled back Topic edits made since that backup:"
)
add_bullets(doc, [
    "Settings: Topics, Job Topics, filters, personas, company context, message templates, negative topics, "
    "the imported Target Accounts list and its score threshold (v0.28.0, see 6.11), API key (opt-in per "
    "export, for deliberately sharing a spend-capped trial key). Import replaces "
    "wholesale - this is configuration a person deliberately set.",
    "Leads: the full lead dataset, plus the generic Sales Mentor and Customer Voice conversation histories "
    "(previously not backed up anywhere at all, despite being genuinely irreplaceable). Import merges, "
    "never replaces - a lead or chat history already present locally is left exactly as-is; only what's "
    "genuinely missing locally gets restored, so importing an older backup can never discard newer local "
    "activity.",
])
doc.add_paragraph(
    "Both files download automatically before every scan, so a scan-time failure never loses accumulated "
    "data. Downloads land in /backup (v0.27.0), not the project root - CSV exports land in /exports. "
    "“Clear Results” (side panel) is display-only - it clears that panel's own list, never "
    "chrome.storage.local; no action in the app deletes a saved lead except a person's own explicit "
    "per-lead status change."
)
p = doc.add_paragraph()
r = p.add_run("Empty-install warning (v0.27.1): "); r.bold = True
p.add_run(
    "Chrome ties an unpacked extension's storage to its install location, not its code - moving or "
    "reinstalling from a different path starts genuinely blank even though the old data still exists in a "
    "backup file, and no API lets a new install read an old one's storage to auto-migrate it. The side "
    "panel shows a banner whenever it finds zero Topics/Job Topics and zero leads, pointing directly at "
    "Import Settings/Import Leads instead of leaving an unexplained blank slate - worded to also make sense "
    "for a genuinely new install."
)

doc.add_heading("6.9 Help", level=3)
doc.add_paragraph(
    "A dedicated Help page (its own tab, opened from a “Help ↗” button next to Dashboard/"
    "Advisors/Settings): a curated set of Q&A entries covering every feature above, plus a free-text search "
    "box. Search is deliberately not an AI feature — no API key needed, instant, and answers are fixed/"
    "reviewed rather than generated. Matching is token-overlap across each entry's question and synonym "
    "keywords (so “enable a Topic” and “disable a search topic” surface the same entry) "
    "with light typo tolerance (edit-distance on individual words). Each result is a collapsible card."
)

doc.add_heading("6.10 Activity Log", level=3)
doc.add_paragraph(
    "A dedicated page (v0.26.0, its own tab, opened from an “Activity Log ↗” button next to "
    "Help) recording every meaningful User action and automatic Extension action, with old/new values where "
    "applicable and every error - prompted by two incidents in one session that were hard to diagnose "
    "without it: an AND-topic bug only visible via the background service worker's own DevTools console "
    "(unreliable - it clears itself when the worker goes idle), and a data-loss incident reconstructed "
    "after the fact from context clues."
)
add_bullets(doc, [
    "Written directly to chrome.storage.local from wherever each action actually happens - including "
    "inside background.js itself - never inferred from chrome.runtime.sendMessage broadcasts, which are "
    "lost entirely if no page happens to be listening (true of every scan-lifecycle message already, with "
    "no storage-backed fallback before this). The log is therefore complete even for a scan that ran while "
    "every page was closed.",
    "Covers: Topic/Job Topic/Negative Topic add/remove/enable/scope/keyword edits (old→new keyword counts), "
    "scans started/completed/errored, automatic negative-topic filtering (both on discovery and on "
    "re-apply) and automatic company extraction/prioritization (including correlated re-scoring) with "
    "their failures, Export/Import Settings and Leads, lead status/priority/company changes, Bulk Change + "
    "Undo, Prioritize/Re-score/Extract Companies, Apply Negative Filters, accepted Lookalike/Search-Quality "
    "suggestions, settings field edits, and clearing an AI conversation.",
    "Free-text fields log once per real edit (focus → blur, only if changed), not per keystroke - reverting "
    "a field back to its original value logs nothing. The Anthropic API key's actual value is never "
    "logged, only that it changed.",
    "Capped at 2000 entries (oldest dropped first) - no unlimitedStorage permission is declared (5MB real "
    "chrome.storage.local quota) and nothing else in the app guards against quota exhaustion, so the log is "
    "Never manually clearable (v0.26.2) - this is the one place to investigate what happened after "
    "something looks wrong, so no action anywhere deletes it. Stored as one array per calendar day "
    "(activityLog:YYYY-MM-DD) rather than one shared array; anything older than a 90-day retention window "
    "is pruned automatically on every write - a predictable “always the last 90 days” guarantee, "
    "not a raw entry-count cap that could exhaust itself faster during a single unusually active day. Log "
    "data from the earlier flat single-key scheme migrates automatically (bucketed by each entry's own "
    "timestamp) the first time it's read.",
    "Updates live via chrome.storage.onChanged (v0.26.1) - a scan can log many entries over its whole run, "
    "and this page doesn't require a manual reload to see them, the same reactive pattern the Dashboard/"
    "Advisors pages already use for their own storage reads.",
    "Periodic file export to /log (v0.27.0) - since this app never runs anything in the background on its "
    "own, this piggybacks on the existing manual Scan trigger (the same moment Settings/Leads backups "
    "already fire) rather than a chrome.alarms schedule. Each closed day (not today, which is still being "
    "written to) is exported to log/activityLog-YYYY-MM-DD.json exactly once, the first time a scan happens "
    "on or after the next day - a predictable, permission-free approximation of a daily export, not a true "
    "cron.",
])

doc.add_heading("6.11 Target Accounts (v0.28.0, extended v0.28.1/v0.28.2)", level=3)
doc.add_paragraph(
    "A curated, externally-researched list of target companies - one row per company, scored 0-100 for "
    "AI-consulting sales fit (AI_Priority_Score), with a categorical label (Very High, Very High - "
    "Provisional, High, High - Provisional, Out of Scope, Insufficient Evidence), industry, and top AI "
    "initiatives - feeding directly into lead prioritization (6.4) so a scanned lead at a well-researched, "
    "high-priority company doesn't have to wait on/rely purely on the AI's own judgment of an unfamiliar name."
)
add_bullets(doc, [
    "Source & import: maintained externally as an Excel workbook (ChatGPT-researched, updated roughly every "
    "few months). The Settings page's “Import Target Accounts” button (same file-picker pattern "
    "as Import Settings/Import Leads, 6.8) reads the .xlsx file directly - no conversion step. "
    "code/xlsx-lite.js is a small, dependency-free in-browser reader (v0.28.1) purpose-built for this one "
    "sheet: it unzips the workbook (native DecompressionStream) and reads its Companies sheet's XML (native "
    "DOMParser) rather than bundling a third-party xlsx library into an extension that already holds "
    "LinkedIn/Anthropic host permissions. code/convert_target_accounts.py (the original v0.28.0 approach, "
    "converting to exports/target-accounts.json first) still works and is kept as an optional offline/CLI "
    "path, but isn't part of the normal workflow anymore. Matched against a lead's company field via the "
    "existing normalizeCompanyName() (already used for company grouping elsewhere), so exact legal-suffix/"
    "punctuation differences don't block a match.",
    "Deterministic Priority 1 - a scanned “New” lead with no priority yet, whose company matches "
    "a target account labeled Very High or High (never a Provisional/Insufficient Evidence/Out of Scope "
    "one) at or above a configurable score threshold (Settings, default 70), is automatically set to "
    "Priority 1 during the scan's post-processing - before the batch AI prioritization pass runs, so it's "
    "never double-scored. The tooltip on the Dashboard's priority pill names the match, its score, and its "
    "top AI initiative. The lead also gets a targetAccountMatch: true flag, so it stays distinguishable "
    "from a Mentor-scored lead even though both render the same way.",
    "Soft signal otherwise - a match that doesn't clear the bar above (Provisional, or below threshold) is "
    "still passed into the same batch AI prioritization call (6.4) as context, so the Sales Mentor weighs it "
    "alongside its usual judgment rather than ignoring it outright; when it materially affects the call, the "
    "Mentor's own reason text (also shown as the pill's tooltip) says so explicitly.",
    "Applies everywhere prioritization runs (v0.28.2) - the same split "
    "(partitionLeadsByTargetAccount in storage.js) runs during a scan's automatic pass and both "
    "Dashboard buttons (6.4), not just at scan time. Concretely: importing a new/updated Target "
    "Accounts list and then clicking Re-score All Priorities immediately re-prioritizes every "
    "eligible existing lead - no new scan needed.",
    "Threshold and the imported list itself live in Settings (6.7) and travel with a Settings export/import "
    "(6.8), so a fresh install or a restored backup doesn't lose them.",
])

doc.add_heading("7. Non-functional requirements", level=2)
add_bullets(doc, [
    "Manual-trigger only — no alarms, no background scanning, ever.",
    "Local-first privacy — all data in chrome.storage.local; the only outbound calls are to linkedin.com "
    "(reading pages already open) and api.anthropic.com (only when AI features are used, with the user's "
    "own key). No server operated by this project.",
    "No remote code — no bundler-fetched or eval'd remote JavaScript (a Chrome Web Store policy "
    "requirement); Dashboard's pie charts are hand-drawn inline SVG rather than a chart library for this "
    "reason.",
    "Self-healing data migrations — schema changes (e.g. the \u201cBlocked\u201d \u2192 \u201cIrrelevant\u201d "
    "rename, missing postedAt/status on old leads) backfill automatically on next read, never requiring a "
    "manual migration step or losing existing data.",
    "Bounded AI calls — every Anthropic fetch has a timeout; a stalled connection now fails with a clear "
    "message inside whatever bound is appropriate (2.5 min conversational, 45s draft, 50s tool execution) "
    "instead of hanging indefinitely.",
])

doc.add_heading("8. Open items / known gaps", level=2)
add_bullets(doc, [
    "LinkedIn's DOM structure for search results is unstable; content-script selectors need occasional "
    "maintenance when scraping breaks.",
    "Negative-topic and Author-title matching is keyword-based, not exact — a common word can over-match, "
    "which is why every automatic status change is reviewable and reversible, never a silent delete.",
    "No CRM push (Salesforce/other) yet — leads live only in this extension.",
    "Company grouping is normalized-string matching, not brand/subsidiary aliasing - it handles minor "
    "legal-suffix variation (“Azqore” vs “Azqore SA”) but has no way to know that "
    "different names are the same company (e.g. “Google Inc.” vs “Alphabet,” "
    "“Facebook” vs “Meta”). A proper alias/merge mechanism is planned as later work, "
    "likely alongside a dedicated Companies view (deferred for now in favor of the in-table grouping above).",
    "Chrome Web Store submission was rejected once (v0.15.0, excessive scripting/tabs permissions that "
    "weren't actually used) and resubmitted after removing them (v0.15.1); awaiting review as of this "
    "writing.",
])

doc.add_heading("9. Version history", level=2)
doc.add_paragraph(
    "See RELEASE_NOTES.md for the full, dated changelog. Current version: 0.28.2."
)

for section in doc.sections:
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)

prd_path = SALESTEAM_DIR + "/PRD.docx"
doc.save(prd_path)
fix_zoom(prd_path)
print("Saved PRD.docx")
