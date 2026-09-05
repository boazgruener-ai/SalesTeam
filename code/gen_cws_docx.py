# -*- coding: utf-8 -*-
import zipfile
import shutil
import re
import os
import docx
from docx.shared import Pt

SALESTEAM_DIR = "C:/Users/41793/OneDrive/Documents/Personal/Claude Code/SalesTeam"


def fix_zoom(path):
    tmp = path + ".tmp"
    shutil.copy(path, tmp)
    with zipfile.ZipFile(tmp, "r") as zin:
        names = zin.namelist()
        settings = zin.read("word/settings.xml").decode("utf-8")
        if "<w:zoom" not in settings:
            m = re.search(r"(<w:settings[^>]*>)", settings)
            if m:
                settings = settings[:m.end()] + '<w:zoom w:percent="100"/>' + settings[m.end():]
        data = {n: zin.read(n) for n in names}
        data["word/settings.xml"] = settings.encode("utf-8")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    os.remove(tmp)


doc = docx.Document()

doc.add_heading("Chrome Web Store listing \u2014 copy-paste reference", level=1)

doc.add_heading("Store listing name", level=2)
doc.add_paragraph("SalesTeam - AI Sales Team for LinkedIn")

doc.add_heading("Summary (max 132 characters)", level=2)
doc.add_paragraph(
    "Scan LinkedIn for leads, auto-filter noise, get AI-prioritized results in a Dashboard, and draft "
    "outreach with an AI sales team."
)

doc.add_heading("Category", level=2)
doc.add_paragraph('Productivity (or "Business tools" if offered as a subcategory)')

doc.add_heading("Language", level=2)
doc.add_paragraph("English")

doc.add_heading("Detailed description", level=2)
doc.add_paragraph(
    "SalesTeam turns LinkedIn's Posts and Jobs search into a prioritized pipeline of leads - then gives you "
    "an AI sales team to act on it."
)

p = doc.add_paragraph()
p.add_run("WHAT IT DOES").bold = True
for item in [
    'Define "Topics" (keyword groups) once, then click "Scan All Topics" to search LinkedIn Posts and Jobs '
    "for all of them in one go.",
    'Define "Negative Topics" too - competitors and recruiter/staffing posts that match the same keywords '
    "but are never real prospects get auto-filtered out, not left for you to skip past manually. Two are "
    "built in; add your own for any other recurring noise.",
    "Results are merged, deduplicated, and ranked, with matched keywords, hiring/freelance signals, and "
    "connection-degree shown on every lead.",
    "A Dashboard tab turns the results into a real pipeline view: pie charts by status, a "
    "sortable/filterable table, per-lead detail pages, CSV export, and safe bulk status changes "
    "(confirmation required, one level of undo).",
    "100% manual-trigger. Nothing runs on a schedule or in the background - every scan starts with you "
    "clicking a button, so it behaves like a careful human, not an automation bot.",
]:
    doc.add_paragraph(item, style="List Bullet")

p = doc.add_paragraph()
p.add_run("YOUR AI SALES TEAM (optional, needs your own Anthropic API key)").bold = True
for item in [
    "Automatic prioritization: after every scan, each new lead is scored 1 (highest) to 5 (lowest) by real "
    "fit and urgency - not just keyword overlap - so you always know what to work on first.",
    "AI-drafted opening messages for any lead - editable, and only ever copied for you to paste and send "
    "yourself. Nothing is auto-sent.",
    'Sales Mentor: an AI agent you can ask "Which lead should I approach first?" or general sales-strategy '
    "questions. It looks up your real lead data only when a question actually needs it.",
    "Customer Voice: an AI agent that roleplays as a realistic buyer, so you can pressure-test a message or "
    "approach before you send it.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "The core scanning, filtering, and Dashboard features work with no AI key at all. The AI Sales Team "
    "features (prioritization, drafting, Sales Mentor, Customer Voice) are entirely optional, and require "
    "you to provide your own Anthropic API key, which is stored only on your device."
)

p = doc.add_paragraph()
p.add_run("PRIVACY").bold = True
doc.add_paragraph(
    "SalesTeam only reads LinkedIn Posts/Jobs search pages you're already viewing, stores everything "
    "locally in your browser, and never sends anything to a server we operate. See the full privacy policy "
    "at: https://claude.ai/code/artifact/727687e7-bf9b-4f4f-b3c1-5ce26b341049"
)

doc.add_heading("Single purpose description (required field)", level=2)
doc.add_paragraph(
    "Scans LinkedIn's Posts and Jobs search results for user-defined topics to surface B2B sales leads, "
    "using the same logged-in LinkedIn session the user would browse with manually \u2014 no separate "
    "credentials or automation account. Also offers optional AI-assisted analysis and message drafting "
    "using the user's own Anthropic API key."
)

doc.add_heading("Permission justifications (required field per permission)", level=2)

perms = [
    ("storage:", "Stores the user's search Topics, filters, scraped leads, and AI conversation history "
     "locally in the browser."),
    ("sidePanel:", "Displays the extension's core UI (search Topics, Negative Topics/filters, scan trigger, "
     "and results) in Chrome's side panel, with links to open the Dashboard, Advisors, Settings, and Help "
     "pages in their own tabs."),
    ("clipboardWrite:", "Lets the user copy an AI-drafted message with one click, to paste into LinkedIn's "
     "own message compose box."),
    ("Host permission - https://www.linkedin.com/*:", "Required to read the Posts and Jobs search-results "
     "pages the user is already viewing, in order to scrape and match leads. The extension does not run on "
     "any other LinkedIn page or any other website."),
    ("Host permission - https://api.anthropic.com/*:", "Required so the optional AI features (message "
     "drafting, Sales Mentor, Customer Voice) can call Anthropic's API directly from the browser, "
     "authenticated with the user's own API key."),
]
for label, body in perms:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    doc.add_paragraph(body)

doc.add_heading("Remote code question", level=2)
doc.add_paragraph(
    "Answer: No. The extension makes data API calls to Anthropic (text in, text out) - it does not fetch "
    "or execute remote JavaScript."
)

doc.add_heading('Data usage disclosure (Chrome\'s "Privacy practices" tab)', level=2)
for item in [
    "Does this extension collect or transmit user data? Yes - lead text and chat content, only when the "
    "optional AI features are used, sent directly to Anthropic's API using the user's own key.",
    "Is this data sold or used for purposes unrelated to the extension's function? No.",
    "Is this data used for advertising? No.",
    "Privacy policy URL: https://claude.ai/code/artifact/727687e7-bf9b-4f4f-b3c1-5ce26b341049",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Visibility", level=2)
doc.add_paragraph(
    'Set to "Unlisted" (not "Public") so it\'s installable only by people you send the link to, without '
    "appearing in Chrome Web Store search."
)

for section in doc.sections:
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)

out_path = SALESTEAM_DIR + "/Chrome Web Store Listing.docx"
doc.save(out_path)
fix_zoom(out_path)
print("Saved Chrome Web Store Listing.docx")
